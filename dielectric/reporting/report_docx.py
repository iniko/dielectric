"""Render a :class:`ReportData` to a Word publication appendix (requires the ``report`` extra)."""

from __future__ import annotations

from pathlib import Path

from .report import ReportData


def _require_docx() -> object:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - exercised only without the optional dep
        raise ImportError(
            "DOCX reports need python-docx; install with `pip install dielectric[report]`."
        ) from exc
    return docx


def new_docx() -> tuple[object, object]:
    """A blank Word document + the docx module (shared by per-report and combined renderers)."""
    docx = _require_docx()
    return docx.Document(), docx  # type: ignore[attr-defined]


def write_report_docx(document: object, report: ReportData, docx: object) -> None:
    """Write one :class:`ReportData` (heading + sections) into an existing Word document."""
    document.add_heading(report.title, level=0)  # type: ignore[attr-defined]
    document.add_paragraph(report.validation_status).runs[0].bold = True  # type: ignore[attr-defined]

    document.add_heading("Methods", level=1)  # type: ignore[attr-defined]
    document.add_paragraph(report.methods)  # type: ignore[attr-defined]

    document.add_heading("Fitted parameters", level=1)  # type: ignore[attr-defined]
    table = document.add_table(rows=1, cols=2)  # type: ignore[attr-defined]
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    for name, value in (*report.parameter_rows, *report.gof_rows):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = value

    document.add_heading("Model selection", level=1)  # type: ignore[attr-defined]
    sel_para = document.add_paragraph(report.selection_table, style="No Spacing")  # type: ignore[attr-defined]
    sel_para.runs[0].font.name = "Courier New"
    for note in report.extra_notes:
        document.add_paragraph(note, style="List Bullet")  # type: ignore[attr-defined]

    for fig_path in report.figure_paths:
        if Path(fig_path).exists():
            document.add_picture(fig_path, width=docx.shared.Inches(6.0))  # type: ignore[attr-defined]

    document.add_heading("References (BibTeX)", level=1)  # type: ignore[attr-defined]
    document.add_paragraph(report.bibtex, style="No Spacing")  # type: ignore[attr-defined]
    document.add_heading("Reproducibility manifest", level=1)  # type: ignore[attr-defined]
    document.add_paragraph(report.manifest_json, style="No Spacing")  # type: ignore[attr-defined]


def render_docx(report: ReportData, path: str | Path) -> None:
    """Write ``report`` as a .docx publication appendix."""
    document, docx = new_docx()
    write_report_docx(document, report, docx)
    document.save(str(path))  # type: ignore[attr-defined]
